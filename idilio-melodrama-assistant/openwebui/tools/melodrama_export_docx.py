"""
title: Melodrama Export a DOCX
author: Idilio
description: Convierte la parte exportable de un guion.md (desde EXPORT-START) a un archivo .docx, para la skill Melodrama Script Intelligence.
required_open_webui_version: 0.5.0
requirements: python-docx
version: 0.1.0
"""

import os
import re

from docx import Document
from docx.shared import Pt
from pydantic import BaseModel, Field

EXPORT_MARKER = '<!-- EXPORT-START -->'


def _slug(show_slug: str) -> str:
    cleaned = re.sub(r'[^a-z0-9-]+', '-', show_slug.strip().lower()).strip('-')
    if not cleaned:
        raise ValueError('show_slug inválido')
    return cleaned


class Tools:
    class Valves(BaseModel):
        BASE_DIR: str = Field(
            default='/app/backend/data/melodrama-guiones',
            description='Directorio base donde vive guiones/<show-slug>/guion.md y donde se escribe el .docx.',
        )

    def __init__(self):
        self.valves = self.Valves()

    def export_to_docx(self, show_slug: str, title: str = '') -> str:
        """
        Genera guiones/<show-slug>/guion.docx a partir del contenido de guion.md que está después del marcador EXPORT-START (nunca del material de desarrollo de arriba del marcador). Úsala solo cuando el libretista pida explícitamente exportar a Word.
        :param show_slug: El identificador del show en minúsculas-con-guiones.
        :param title: Título a mostrar en la primera línea del documento (opcional; por defecto usa el show_slug).
        :return: Confirmación con la ruta del .docx generado, o un mensaje de error si el guion.md no existe o no tiene el marcador EXPORT-START.
        """
        slug = _slug(show_slug)
        md_path = os.path.join(self.valves.BASE_DIR, slug, 'guion.md')
        if not os.path.exists(md_path):
            return f"No existe guion.md para '{show_slug}' -- no hay nada que exportar."

        with open(md_path, encoding='utf-8') as f:
            content = f.read()

        if EXPORT_MARKER not in content:
            return f'guion.md no tiene el marcador {EXPORT_MARKER} -- no se puede exportar.'

        script_text = content.split(EXPORT_MARKER, 1)[1].strip()
        if not script_text:
            return 'La parte exportable del guion.md está vacía todavía -- no hay capítulos que exportar.'

        document = Document()
        document.styles['Normal'].font.size = Pt(11)
        document.add_heading(title or slug.replace('-', ' ').title(), level=0)

        for block in script_text.split('\n\n'):
            block = block.strip()
            if not block:
                continue
            if re.match(r'^CAP[IÍ]TULO\s+\d+', block, re.IGNORECASE):
                document.add_heading(block, level=1)
            elif re.match(r'^\d+\.\s+(INT|EXT)', block, re.IGNORECASE):
                paragraph = document.add_paragraph(block)
                paragraph.runs[0].bold = True
            else:
                document.add_paragraph(block)

        docx_path = os.path.join(self.valves.BASE_DIR, slug, 'guion.docx')
        document.save(docx_path)
        return f'Exportado: {docx_path}'
